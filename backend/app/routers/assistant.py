import io
import os
import re
import uuid
from datetime import date
from typing import List, Optional
from fastapi import APIRouter, Depends, HTTPException, Request, UploadFile, File, Form
from sqlalchemy.orm import Session
import pypdf
import docx

from app.core.database import get_db
from app.core.security import get_current_user, RoleChecker
from app.core.utils import log_audit
from app.core.ai import generate_ai_chat_response
from app.models.models import User, CompanyPolicy, OfficeSetting, Holiday, Attendance
from app.schemas.schemas import (
    CompanyPolicyOut, CompanyPolicyCreate, CompanyPolicyUpdate,
    DocumentExtractResponse, AIChatRequest, AIChatResponse, MessageResponse
)

router = APIRouter(prefix="/assistant", tags=["AI HR Policy Assistant"])

admin_required = RoleChecker(["Admin"])

MAX_DOC_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extracts text from PDF, DOCX, TXT, or MD documents.
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        try:
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            pages_text = []
            for page in reader.pages:
                txt = page.extract_text()
                if txt:
                    pages_text.append(txt.strip())
            extracted = "\n\n".join(pages_text).strip()
            if not extracted:
                raise ValueError("PDF contains no extractable text (it may be a scanned image).")
            return extracted
        except Exception as e:
            if isinstance(e, ValueError):
                raise
            raise ValueError(f"Failed to parse PDF document: {str(e)}")

    elif ext in [".docx", ".doc"]:
        if ext == ".doc":
            raise ValueError("Legacy .doc format is not supported. Please convert to modern .docx or .pdf.")
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        paragraphs.append(row_text)

            extracted = "\n\n".join(paragraphs).strip()
            if not extracted:
                raise ValueError("DOCX document contains no extractable text.")
            return extracted
        except Exception as e:
            if isinstance(e, ValueError):
                raise
            raise ValueError(f"Failed to parse DOCX document: {str(e)}")

    elif ext in [".txt", ".md", ".markdown"]:
        try:
            return file_bytes.decode("utf-8").strip()
        except UnicodeDecodeError:
            try:
                return file_bytes.decode("latin-1").strip()
            except Exception as e:
                raise ValueError(f"Failed to decode text document: {str(e)}")

    else:
        raise ValueError(f"Unsupported file format '{ext}'. Allowed formats: .pdf, .docx, .txt, .md")

def infer_title(filename: str, text: str) -> str:
    """
    Derives an appropriate policy title from first header or filename.
    """
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if lines:
        first_line = re.sub(r"^[#*=\-\s]+", "", lines[0]).strip()
        if 4 <= len(first_line) <= 80 and not first_line.endswith("."):
            return first_line

    name_without_ext = os.path.splitext(filename)[0]
    clean_name = re.sub(r"[_\-]+", " ", name_without_ext).strip()
    return clean_name.title() if clean_name else "Company Workplace Policy"

def infer_category(text: str, filename: str) -> str:
    """
    Heuristically infers the policy category based on keywords.
    """
    corpus = (filename + " " + text[:2500]).lower()

    if any(k in corpus for k in ["leave", "vacation", "sick leave", "casual leave", "maternity", "paternity", "time off", "pto", "bereavement", "holiday"]):
        return "Leaves"
    if any(k in corpus for k in ["attendance", "working hour", "shift", "overtime", "clock in", "clock out", "punch", "wfh", "remote work", "timing", "punctuality"]):
        return "Attendance"
    if any(k in corpus for k in ["reimbursement", "travel", "expense", "perk", "allowance", "medical", "insurance", "bonus", "salary", "provident", "gratuity", "benefit", "claim"]):
        return "Benefits"
    if any(k in corpus for k in ["harassment", "conduct", "ethics", "posh", "discipline", "confidentiality", "nda", "dress code", "acceptable use", "behavior", "grievance"]):
        return "Code of Conduct"

    return "General"

# --- Company Policy Knowledge Base Management ---

@router.get("/policies", response_model=List[CompanyPolicyOut])
def get_company_policies(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    query = db.query(CompanyPolicy).filter(
        CompanyPolicy.organization_id == current_user.organization_id
    )
    if current_user.role == "Employee":
        query = query.filter(CompanyPolicy.is_published == True)
    return query.order_by(CompanyPolicy.category.asc(), CompanyPolicy.title.asc()).all()

@router.post("/policies", response_model=CompanyPolicyOut)
def create_company_policy(
    request: Request,
    payload: CompanyPolicyCreate,
    db: Session = Depends(get_db),
    admin_user: User = Depends(admin_required)
):
    new_policy = CompanyPolicy(
        organization_id=admin_user.organization_id,
        title=payload.title.strip(),
        category=payload.category.strip(),
        content=payload.content.strip(),
        is_published=payload.is_published
    )
    db.add(new_policy)
    db.commit()
    db.refresh(new_policy)

    log_audit(db, admin_user.id, "POLICY_CREATED", request.client.host if request.client else None, f"Title: {payload.title}", organization_id=admin_user.organization_id)
    return new_policy

@router.put("/policies/{policy_id}", response_model=CompanyPolicyOut)
def update_company_policy(
    policy_id: str,
    request: Request,
    payload: CompanyPolicyUpdate,
    db: Session = Depends(get_db),
    admin_user: User = Depends(admin_required)
):
    try:
        uuid_id = uuid.UUID(policy_id) if isinstance(policy_id, str) else policy_id
    except (ValueError, TypeError):
        raise HTTPException(status_code=400, detail="Invalid policy ID format")

    policy = db.query(CompanyPolicy).filter(
        CompanyPolicy.id == uuid_id,
        CompanyPolicy.organization_id == admin_user.organization_id
    ).first()
    if not policy:
        raise HTTPException(status_code=404, detail="Company policy not found")

    update_dict = payload.model_dump(exclude_unset=True)
    for k, v in update_dict.items():
        setattr(policy, k, v.strip() if isinstance(v, str) else v)

    db.commit()
    db.refresh(policy)

    log_audit(db, admin_user.id, "POLICY_UPDATED", request.client.host if request.client else None, f"Policy ID: {policy_id}", organization_id=admin_user.organization_id)
    return policy

@router.delete("/policies/{policy_id}", response_model=MessageResponse)
def delete_company_policy(
    policy_id: str,
    request: Request,
    db: Session = Depends(get_db),
    admin_user: User = Depends(admin_required)
):
    try:
        uuid_id = uuid.UUID(policy_id) if isinstance(policy_id, str) else policy_id
    except (ValueError, TypeError):
        raise HTTPException(status_code=400, detail="Invalid policy ID format")

    policy = db.query(CompanyPolicy).filter(
        CompanyPolicy.id == uuid_id,
        CompanyPolicy.organization_id == admin_user.organization_id
    ).first()
    if not policy:
        raise HTTPException(status_code=404, detail="Company policy not found")

    db.delete(policy)
    db.commit()

    log_audit(db, admin_user.id, "POLICY_DELETED", request.client.host if request.client else None, f"Policy ID: {policy_id}", organization_id=admin_user.organization_id)
    return {"message": "Policy deleted successfully"}

# --- Document Extraction & Direct File Ingestion for RAG ---

@router.post("/policies/extract-document", response_model=DocumentExtractResponse)
def extract_policy_document(
    file: UploadFile = File(...),
    admin_user: User = Depends(admin_required)
):
    """
    Extracts text and metadata from an uploaded .pdf, .docx, .txt, or .md file
    so HR can preview, refine, and edit before saving to the AI Knowledge Base.
    """
    # Sync on purpose: the parsing/disk/DB work below is blocking, so it
    # belongs in FastAPI's threadpool rather than on the event loop.
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file selected for upload")

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in [".pdf", ".docx", ".txt", ".md", ".markdown"]:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file format '{ext}'. Allowed formats: .pdf, .docx, .txt, .md"
        )

    file_bytes = file.file.read()
    if len(file_bytes) > MAX_DOC_SIZE_BYTES:
        raise HTTPException(status_code=400, detail="Document exceeds maximum allowed size of 10 MB")
    if len(file_bytes) == 0:
        raise HTTPException(status_code=400, detail="Uploaded document is empty")

    try:
        content = extract_text_from_file(file_bytes, file.filename)
    except ValueError as ve:
        raise HTTPException(status_code=400, detail=str(ve))

    title = infer_title(file.filename, content)
    suggested_category = infer_category(content, file.filename)

    return DocumentExtractResponse(
        title=title[:200],
        suggested_category=suggested_category,
        content=content,
        filename=file.filename,
        character_count=len(content)
    )

@router.post("/policies/upload-file", response_model=CompanyPolicyOut)
def upload_policy_file(
    request: Request,
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    category: Optional[str] = Form(None),
    is_published: bool = Form(True),
    db: Session = Depends(get_db),
    admin_user: User = Depends(admin_required)
):
    """
    Directly uploads and parses a document file, saving it as a Company Policy
    in the tenant's AI Knowledge Base in one step.
    """
    # Sync on purpose: the parsing/disk/DB work below is blocking, so it
    # belongs in FastAPI's threadpool rather than on the event loop.
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file selected for upload")

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in [".pdf", ".docx", ".txt", ".md", ".markdown"]:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file format '{ext}'. Allowed formats: .pdf, .docx, .txt, .md"
        )

    file_bytes = file.file.read()
    if len(file_bytes) > MAX_DOC_SIZE_BYTES:
        raise HTTPException(status_code=400, detail="Document exceeds maximum allowed size of 10 MB")
    if len(file_bytes) == 0:
        raise HTTPException(status_code=400, detail="Uploaded document is empty")

    try:
        content = extract_text_from_file(file_bytes, file.filename)
    except ValueError as ve:
        raise HTTPException(status_code=400, detail=str(ve))

    final_title = (title.strip() if title and title.strip() else infer_title(file.filename, content))[:200]
    final_category = (category.strip() if category and category.strip() else infer_category(content, file.filename))[:50]

    new_policy = CompanyPolicy(
        organization_id=admin_user.organization_id,
        title=final_title,
        category=final_category,
        content=content,
        is_published=is_published
    )
    db.add(new_policy)
    db.commit()
    db.refresh(new_policy)

    log_audit(
        db,
        admin_user.id,
        "POLICY_UPLOADED",
        request.client.host if request.client else None,
        f"File: {file.filename}, Title: {final_title}",
        organization_id=admin_user.organization_id
    )
    return new_policy

# --- AI Assistant Q&A Chatbot Endpoint ---

@router.post("/chat", response_model=AIChatResponse)
def chat_with_hr_assistant(
    payload: AIChatRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    org_id = current_user.organization_id
    company_name = current_user.organization.name if current_user.organization else "AuraWork"

    # 1. Fetch Employee Live Profile Context
    profile = current_user.profile
    dept_name = profile.department.name if profile and profile.department else "General"
    emp_name = f"{profile.first_name} {profile.last_name}".strip() if profile else current_user.email
    
    today = date.today()
    today_attendance = db.query(Attendance).filter(
        Attendance.organization_id == org_id,
        Attendance.user_id == current_user.id,
        Attendance.date == today
    ).first()
    today_status = today_attendance.status if today_attendance else "Not Clocked In"

    employee_context = {
        "name": emp_name,
        "email": current_user.email,
        "role": current_user.role,
        "department": dept_name,
        "designation": profile.designation if profile and profile.designation else "Employee",
        "leave_balance_casual": profile.leave_balance_casual if profile else 0,
        "leave_balance_sick": profile.leave_balance_sick if profile else 0,
        "leave_balance_paid": profile.leave_balance_paid if profile else 0,
        "wfh_enabled": profile.wfh_enabled if profile else False,
        "today_status": today_status
    }

    # 2. Fetch Tenant Policies (only published)
    policies_query = db.query(CompanyPolicy).filter(
        CompanyPolicy.organization_id == org_id,
        CompanyPolicy.is_published == True
    ).all()
    policies_data = [
        {"title": p.title, "category": p.category, "content": p.content}
        for p in policies_query
    ]

    # 3. Fetch Tenant Office Settings
    setting = db.query(OfficeSetting).filter(OfficeSetting.organization_id == org_id).first()
    office_settings_data = None
    if setting:
        office_settings_data = {
            "office_start_time": setting.office_start_time.strftime("%H:%M"),
            "office_end_time": setting.office_end_time.strftime("%H:%M"),
            "lunch_break_hours": setting.lunch_break_hours,
            "required_working_hours": setting.required_working_hours,
            "weekends": setting.weekends,
            "timezone": setting.timezone
        }

    # 4. Fetch Holidays
    holidays_query = db.query(Holiday).filter(
        Holiday.organization_id == org_id,
        Holiday.date >= today
    ).order_by(Holiday.date.asc()).limit(10).all()
    holidays_data = [
        {"name": h.name, "date": str(h.date), "description": h.description}
        for h in holidays_query
    ]

    # 5. Generate AI Chat Response
    chat_history_dicts = [
        {"role": m.role, "content": m.content}
        for m in (payload.history or [])
    ]

    result = generate_ai_chat_response(
        user_message=payload.message.strip(),
        company_name=company_name,
        employee_context=employee_context,
        policies=policies_data,
        office_settings=office_settings_data,
        holidays=holidays_data,
        chat_history=chat_history_dicts
    )

    return AIChatResponse(
        reply=result["reply"],
        sources=result.get("sources", [])
    )
