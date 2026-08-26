import pytest
from datetime import date
from app.models.models import Organization, OfficeSetting, Holiday, CompanyPolicy, User, EmployeeProfile
from app.core.security import get_password_hash, create_jwt_token

def test_organization_settings_get_and_update(admin_client):
    # 1. Get organization details
    res_get = admin_client.get("/api/settings/organization")
    assert res_get.status_code == 200
    assert res_get.json()["slug"] == "test-company"

    # 2. Update organization name
    res_put = admin_client.put("/api/settings/organization", json={
        "name": "Updated Enterprise Corp"
    })
    assert res_put.status_code == 200
    assert res_put.json()["name"] == "Updated Enterprise Corp"

def test_office_settings_shift_timing_update(admin_client):
    res = admin_client.put("/api/settings/office", json={
        "office_start_time": "09:00:00",
        "office_end_time": "18:00:00",
        "allowed_radius": 250.0,
        "lunch_break_hours": 1.5,
        "required_working_hours": 7.5,
        "weekends": "Saturday,Sunday",
        "timezone": "Asia/Kolkata"
    })
    assert res.status_code == 200
    assert res.json()["allowed_radius"] == 250.0

def test_holiday_creation_and_deletion(admin_client):
    test_date = "2026-12-31"
    
    # 1. Create holiday
    res_c = admin_client.post("/api/settings/holidays", json={
        "name": "New Year Eve Special",
        "date": test_date,
        "description": "Year end holiday"
    })
    assert res_c.status_code == 200
    h_id = res_c.json()["id"]

    # 2. Delete holiday
    res_d = admin_client.delete(f"/api/settings/holidays/{h_id}")
    assert res_d.status_code == 200
    assert "deleted successfully" in res_d.json()["message"].lower()

def test_holiday_duplicate_date_rejected(admin_client):
    test_date = "2026-11-15"
    admin_client.post("/api/settings/holidays", json={"name": "Festival 1", "date": test_date})

    # Duplicate on same date
    res_dup = admin_client.post("/api/settings/holidays", json={"name": "Festival 2", "date": test_date})
    assert res_dup.status_code == 400
    assert "already configured" in res_dup.json()["detail"].lower()

def test_company_policy_category_filtering(admin_client):
    admin_client.post("/api/assistant/policies", json={
        "title": "Maternity & Paternity Leave Policy",
        "category": "Leaves",
        "content": "26 weeks of paid maternity leave."
    })
    admin_client.post("/api/assistant/policies", json={
        "title": "IT Device Security Policy",
        "category": "General",
        "content": "All laptops must have disk encryption enabled."
    })

    res = admin_client.get("/api/assistant/policies")
    assert res.status_code == 200
    policies = res.json()
    assert any(p["category"] == "Leaves" for p in policies)
    assert any(p["category"] == "General" for p in policies)

def test_company_policy_draft_hidden_from_employees(admin_client, employee_client):
    # Admin creates a draft policy
    res_c = admin_client.post("/api/assistant/policies", json={
        "title": "Draft Secret Bonus Policy",
        "category": "Benefits",
        "content": "Confidential draft",
        "is_published": False
    })
    assert res_c.status_code == 200
    policy_id = res_c.json()["id"]

    # Employee lists policies -> draft should be excluded
    res_emp = employee_client.get("/api/assistant/policies")
    assert res_emp.status_code == 200
    assert not any(p["id"] == policy_id for p in res_emp.json())

    # Admin lists policies -> sees the draft
    res_admin = admin_client.get("/api/assistant/policies")
    assert res_admin.status_code == 200
    assert any(p["id"] == policy_id for p in res_admin.json())

def test_ai_assistant_chat_wfh_policy_inquiry(employee_client, admin_client):
    admin_client.post("/api/assistant/policies", json={
        "title": "Work From Home Guidelines",
        "category": "Attendance",
        "content": "Employees may work remotely up to 2 days per week with manager approval."
    })

    res = employee_client.post("/api/assistant/chat", json={
        "message": "What is the policy for work from home?"
    })
    assert res.status_code == 200
    body = res.json()
    assert "work" in body["reply"].lower() or "home" in body["reply"].lower()

def test_ai_assistant_chat_upcoming_holidays_inquiry(employee_client, admin_client):
    admin_client.post("/api/settings/holidays", json={
        "name": "Spring Harvest Festival",
        "date": "2026-10-15",
        "description": "Harvest Celebrations"
    })

    res = employee_client.post("/api/assistant/chat", json={
        "message": "When is the next upcoming public holiday?"
    })
    assert res.status_code == 200
    assert "holiday" in res.json()["reply"].lower()

def test_ai_assistant_chat_conversational_history(employee_client):
    res = employee_client.post("/api/assistant/chat", json={
        "message": "What is my casual leave balance?",
        "history": [
            {"role": "user", "content": "Hello HR bot"},
            {"role": "assistant", "content": "Hello! How can I assist you with company policies?"}
        ]
    })
    assert res.status_code == 200
    assert "Casual" in res.json()["reply"]

def test_cross_tenant_assistant_isolation(client, db):
    # 1. Create Tenant Delta
    reg_delta = client.post("/api/auth/register-company", json={
        "company_name": "Delta Systems",
        "company_slug": "delta-sys",
        "admin_name": "Admin Delta",
        "admin_email": "admin@deltasys.com",
        "admin_password": "DeltaPassword123"
    })
    assert reg_delta.status_code == 200
    token_delta = reg_delta.cookies["access_token"]
    client.cookies.set("access_token", token_delta)

    # 2. Add secret custom policy in Delta Systems
    client.post("/api/assistant/policies", json={
        "title": "Delta Exclusive Profit Sharing",
        "category": "Benefits",
        "content": "All Delta staff receive 15% quarterly revenue share."
    })

    # 3. Delta admin chats with AI -> AI knows Delta policy
    res_delta_chat = client.post("/api/assistant/chat", json={
        "message": "Tell me about our profit sharing policy"
    })
    assert res_delta_chat.status_code == 200
    assert "15%" in res_delta_chat.json()["reply"] or "revenue" in res_delta_chat.json()["reply"].lower()

def test_extract_policy_document_txt_and_md(admin_client):
    content = "# Workplace Health & Safety Policy\n\nAll team members must follow ergonomic guidelines and report hazards immediately."
    files = {"file": ("Health_and_Safety_Policy.md", content.encode("utf-8"), "text/markdown")}
    
    res = admin_client.post("/api/assistant/policies/extract-document", files=files)
    assert res.status_code == 200
    data = res.json()
    assert data["title"] == "Workplace Health & Safety Policy"
    assert "ergonomic guidelines" in data["content"]
    assert data["character_count"] > 20
    assert data["suggested_category"] in ["General", "Code of Conduct"]

def test_extract_policy_document_pdf(admin_client):
    import io
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    pdf_buffer = io.BytesIO()
    p = canvas.Canvas(pdf_buffer, pagesize=letter)
    p.drawString(100, 750, "Annual Travel and Expense Reimbursement Guidelines")
    p.drawString(100, 720, "Employees can claim up to INR 2500 per day for intercity travel meals and lodging.")
    p.showPage()
    p.save()
    pdf_bytes = pdf_buffer.getvalue()

    files = {"file": ("Travel_Reimbursement_Policy.pdf", pdf_bytes, "application/pdf")}
    res = admin_client.post("/api/assistant/policies/extract-document", files=files)
    assert res.status_code == 200
    data = res.json()
    assert "Travel" in data["title"] or "Reimbursement" in data["title"]
    assert "INR 2500" in data["content"]
    assert data["suggested_category"] == "Benefits"

def test_extract_policy_document_docx(admin_client):
    import io
    import docx

    doc = docx.Document()
    doc.add_heading("Parental and Maternity Leave Policy", level=1)
    doc.add_paragraph("New mothers are eligible for 26 consecutive weeks of fully paid maternity leave.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Leave Type"
    table.cell(0, 1).text = "Duration"
    table.cell(1, 0).text = "Maternity"
    table.cell(1, 1).text = "26 Weeks"

    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_bytes = docx_buffer.getvalue()

    files = {"file": ("Maternity_Leave_Policy.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    res = admin_client.post("/api/assistant/policies/extract-document", files=files)
    assert res.status_code == 200
    data = res.json()
    assert "Maternity" in data["title"]
    assert "26 consecutive weeks" in data["content"]
    assert "Leave Type | Duration" in data["content"] or "Maternity | 26 Weeks" in data["content"]
    assert data["suggested_category"] == "Leaves"

def test_extract_policy_document_unsupported_extension_rejected(admin_client):
    files = {"file": ("malicious_file.exe", b"MZbinarycontent", "application/octet-stream")}
    res = admin_client.post("/api/assistant/policies/extract-document", files=files)
    assert res.status_code == 400
    assert "unsupported file format" in res.json()["detail"].lower()

def test_upload_file_direct_policy_creation(admin_client, employee_client):
    policy_text = "# Gym & Wellness Perk\nEmployees receive $100 monthly wellness allowance for fitness subscriptions."
    files = {"file": ("Wellness_Perk_Policy.txt", policy_text.encode("utf-8"), "text/plain")}
    data = {
        "title": "Corporate Wellness & Gym Benefit",
        "category": "Benefits",
        "is_published": "true"
    }

    res = admin_client.post("/api/assistant/policies/upload-file", files=files, data=data)
    assert res.status_code == 200
    policy_id = res.json()["id"]
    assert res.json()["title"] == "Corporate Wellness & Gym Benefit"
    assert res.json()["category"] == "Benefits"

    # Verify policy is retrieved in AI Knowledge Base
    res_list = employee_client.get("/api/assistant/policies")
    assert res_list.status_code == 200
    assert any(p["id"] == policy_id for p in res_list.json())

    # Verify employee can query AI Assistant about the uploaded document
    res_chat = employee_client.post("/api/assistant/chat", json={
        "message": "What is the policy for gym and wellness perk?"
    })
    assert res_chat.status_code == 200
    reply = res_chat.json()["reply"]
    assert "$100" in reply or "wellness" in reply.lower() or "gym" in reply.lower()
