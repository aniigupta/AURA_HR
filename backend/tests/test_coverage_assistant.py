"""
Coverage for the AI knowledge-base ingestion pipeline: every accepted document
format, every rejection, and the title/category inference heuristics.

Document ingestion is the highest-risk untested surface in the assistant: the
extracted text goes straight into the model prompt, so a parser that fails
open, or one that accepts a format it cannot safely read, becomes a
prompt-injection vector rather than a parsing bug.
"""

import io
import uuid

import docx
import pypdf
import pytest
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from app.routers.assistant import (
    extract_text_from_file,
    infer_category,
    infer_title,
)

MISSING_UUID = "00000000-0000-0000-0000-000000000000"


def make_pdf(lines: list[str]) -> bytes:
    buf = io.BytesIO()
    pdf = canvas.Canvas(buf, pagesize=letter)
    y = 720
    for line in lines:
        pdf.drawString(72, y, line)
        y -= 18
    pdf.save()
    return buf.getvalue()


def make_empty_pdf() -> bytes:
    """A structurally valid PDF with no text layer — what a scan looks like."""
    buf = io.BytesIO()
    canvas.Canvas(buf, pagesize=letter).save()
    return buf.getvalue()


def make_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    document = docx.Document()
    for para in paragraphs:
        document.add_paragraph(para)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, value in enumerate(row):
                table.cell(r, c).text = value
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------
# extract_text_from_file — one test per format and per failure
# --------------------------------------------------------------------------


def test_file_030_pdf_text_is_extracted_page_by_page():
    """A normal text PDF yields its content."""
    text = extract_text_from_file(make_pdf(["Notice Period Policy", "Standard notice is 30 days."]), "policy.pdf")
    assert "Notice Period Policy" in text
    assert "30 days" in text


def test_file_016_a_scanned_pdf_is_rejected_with_a_useful_reason():
    """FILE-016 — 'no text layer' must be distinguishable from 'corrupt file'."""
    with pytest.raises(ValueError, match="no extractable text"):
        extract_text_from_file(make_empty_pdf(), "scan.pdf")


def test_file_031_a_corrupt_pdf_reports_a_parse_failure():
    """Random bytes named .pdf are a parse error, not a crash."""
    with pytest.raises(ValueError, match="Failed to parse PDF document"):
        extract_text_from_file(b"%PDF-1.4 this is not really a pdf", "broken.pdf")


def test_file_032_docx_paragraphs_and_tables_are_both_extracted():
    """Policy documents routinely carry entitlements in tables, not prose."""
    text = extract_text_from_file(
        make_docx(["Leave Policy"], [["Type", "Days"], ["Casual", "12"]]), "leave.docx"
    )
    assert "Leave Policy" in text
    assert "Casual | 12" in text


def test_file_033_an_empty_docx_is_rejected():
    """A document with no readable text has nothing to add to the prompt."""
    with pytest.raises(ValueError, match="no extractable text"):
        extract_text_from_file(make_docx([]), "empty.docx")


def test_file_018_a_corrupt_docx_reports_a_parse_failure():
    """FILE-018 — a 400, not a 500."""
    with pytest.raises(ValueError, match="Failed to parse DOCX document"):
        extract_text_from_file(b"PK\x03\x04 not really a docx", "broken.docx")


def test_file_017_legacy_doc_is_refused_with_guidance():
    """FILE-017 — tell the user what to do, not just that it failed."""
    with pytest.raises(ValueError, match="Legacy .doc format is not supported"):
        extract_text_from_file(b"\xd0\xcf\x11\xe0", "old.doc")


@pytest.mark.parametrize("name", ["policy.txt", "policy.md", "policy.markdown"])
def test_file_034_plain_text_formats_are_read_directly(name):
    """TXT and Markdown need no parser."""
    assert extract_text_from_file(b"  Standard notice is 30 days.  ", name) == "Standard notice is 30 days."


def test_file_019_non_utf8_text_falls_back_to_latin1():
    """FILE-019 — a Windows-1252 export must not fail the upload."""
    assert "café" in extract_text_from_file("café policy".encode("latin-1"), "policy.txt")


def test_file_035_an_unsupported_extension_names_the_allowed_set():
    """The error tells the user which formats would have worked."""
    with pytest.raises(ValueError, match=r"Allowed formats: \.pdf, \.docx, \.txt, \.md"):
        extract_text_from_file(b"data", "spreadsheet.xlsx")


# --------------------------------------------------------------------------
# infer_title / infer_category — the heuristics
# --------------------------------------------------------------------------


def test_file_036_a_leading_heading_becomes_the_title():
    """A markdown or underlined heading is the best title available."""
    assert infer_title("upload.pdf", "### Notice Period Policy\nBody text") == "Notice Period Policy"


@pytest.mark.parametrize(
    "text",
    ["A", "This first line is a full sentence and ends with a period.", "x" * 200],
    ids=["too-short", "sentence", "too-long"],
)
def test_file_037_an_unsuitable_first_line_falls_back_to_the_filename(text):
    """Titles come from headings, not from arbitrary opening prose."""
    assert infer_title("expense_reimbursement-policy.pdf", text) == "Expense Reimbursement Policy"


def test_file_038_a_blank_document_still_gets_a_title():
    """
    Neither a heading nor a usable filename still yields something readable.

    "___.pdf" is the case that reaches the constant: the separators are
    stripped, leaving nothing. Note that a filename of ".pdf" does NOT — POSIX
    reads a leading dot as a hidden file rather than an extension, so
    splitext leaves the whole string as the stem and the title becomes ".Pdf".
    Ugly, but harmless, and worth pinning so a future splitext change is visible.
    """
    assert infer_title("___.pdf", "") == "Company Workplace Policy"
    assert infer_title("---.pdf", "") == "Company Workplace Policy"
    assert infer_title(".pdf", "") == ".Pdf"


@pytest.mark.parametrize(
    "corpus,expected",
    [
        ("Casual leave and maternity entitlements", "Leaves"),
        ("Clock in and clock out punctuality rules", "Attendance"),
        ("Travel expense reimbursement and medical insurance", "Benefits"),
        ("POSH harassment and confidentiality obligations", "Code of Conduct"),
        ("General office stationery ordering process", "General"),
    ],
    ids=["leaves", "attendance", "benefits", "conduct", "general"],
)
def test_file_039_category_inference_covers_every_bucket(corpus, expected):
    """AI-034 — each of the five categories is reachable, including the default."""
    assert infer_category(corpus, "document.pdf") == expected


def test_file_040_the_filename_contributes_to_the_category():
    """A bare filename is enough when the body is uninformative."""
    assert infer_category("Refer to the handbook.", "leave-policy.pdf") == "Leaves"


# --------------------------------------------------------------------------
# The upload endpoints
# --------------------------------------------------------------------------


@pytest.mark.parametrize("endpoint", ["/api/assistant/policies/extract-document", "/api/assistant/policies/upload-file"])
def test_file_041_both_ingest_routes_reject_the_same_inputs(admin_client, endpoint):
    """Preview and save-in-one-step must apply identical rules."""
    unsupported = admin_client.post(endpoint, files={"file": ("data.xlsx", b"payload", "application/vnd.ms-excel")})
    assert unsupported.status_code == 400
    assert "Unsupported file format" in unsupported.json()["detail"]

    empty = admin_client.post(endpoint, files={"file": ("empty.txt", b"", "text/plain")})
    assert empty.status_code == 400
    assert empty.json()["detail"] == "Uploaded document is empty"

    oversized = admin_client.post(
        endpoint, files={"file": ("big.txt", b"x" * (10 * 1024 * 1024 + 1), "text/plain")}
    )
    assert oversized.status_code == 400
    assert "10 MB" in oversized.json()["detail"]

    unparseable = admin_client.post(endpoint, files={"file": ("scan.pdf", make_empty_pdf(), "application/pdf")})
    assert unparseable.status_code == 400
    assert "no extractable text" in unparseable.json()["detail"]


def test_file_042_extract_returns_a_preview_without_saving(admin_client):
    """The preview route is read-only: HR edits before anything is stored."""
    before = len(admin_client.get("/api/assistant/policies").json())

    res = admin_client.post(
        "/api/assistant/policies/extract-document",
        files={"file": ("leave_policy.pdf", make_pdf(["### Annual Leave Policy", "Casual leave is 12 days."]), "application/pdf")},
    )
    assert res.status_code == 200
    body = res.json()
    assert body["title"] == "Annual Leave Policy"
    assert body["suggested_category"] == "Leaves"
    assert body["character_count"] == len(body["content"])
    assert body["filename"] == "leave_policy.pdf"

    assert len(admin_client.get("/api/assistant/policies").json()) == before


def test_file_043_upload_saves_with_inferred_metadata(admin_client):
    """With no title or category supplied, both are inferred and persisted."""
    res = admin_client.post(
        "/api/assistant/policies/upload-file",
        files={"file": ("wfh_policy.txt", b"### Remote Work Guidelines\nWFH requires prior approval.", "text/plain")},
    )
    assert res.status_code == 200
    body = res.json()
    assert body["title"] == "Remote Work Guidelines"
    assert body["category"] == "Attendance"
    assert body["is_published"] is True


def test_file_044_explicit_title_and_category_override_inference(admin_client):
    """HR's own labelling wins over the heuristic."""
    res = admin_client.post(
        "/api/assistant/policies/upload-file",
        files={"file": ("notes.txt", b"Some leave related content here.", "text/plain")},
        data={"title": "  Curated Title  ", "category": "  Benefits  ", "is_published": "false"},
    )
    assert res.status_code == 200
    body = res.json()
    assert body["title"] == "Curated Title"
    assert body["category"] == "Benefits"
    assert body["is_published"] is False


# --------------------------------------------------------------------------
# Policy CRUD — malformed ids and cross-tenant misses
# --------------------------------------------------------------------------


@pytest.mark.parametrize("method", ["put", "delete"])
def test_ai_035_a_malformed_policy_id_is_a_400(admin_client, method):
    """AI-035 — the route parses the UUID itself and says so clearly."""
    kwargs = {"json": {"title": "X"}} if method == "put" else {}
    res = getattr(admin_client, method)("/api/assistant/policies/not-a-uuid", **kwargs)
    assert res.status_code == 400
    assert res.json()["detail"] == "Invalid policy ID format"


@pytest.mark.parametrize("method", ["put", "delete"])
def test_ai_036_an_unknown_policy_id_is_a_404(admin_client, method):
    """A well-formed id belonging to no policy in this tenant is a 404."""
    kwargs = {"json": {"title": "X"}} if method == "put" else {}
    res = getattr(admin_client, method)(f"/api/assistant/policies/{MISSING_UUID}", **kwargs)
    assert res.status_code == 404
    assert res.json()["detail"] == "Company policy not found"


def test_ai_029_policy_crud_round_trips(admin_client):
    """AI-029 — create, update and delete all take effect."""
    created = admin_client.post(
        "/api/assistant/policies",
        json={"title": "  Spaced Title  ", "category": "  General  ", "content": "  Body  ", "is_published": True},
    )
    assert created.status_code == 200
    policy_id = created.json()["id"]
    assert created.json()["title"] == "Spaced Title"  # trimmed on the way in

    updated = admin_client.put(
        f"/api/assistant/policies/{policy_id}",
        json={"content": "  Revised body  ", "is_published": False},
    )
    assert updated.status_code == 200
    assert updated.json()["content"] == "Revised body"
    assert updated.json()["is_published"] is False
    assert updated.json()["title"] == "Spaced Title"  # untouched by a partial update

    assert admin_client.delete(f"/api/assistant/policies/{policy_id}").status_code == 200
    assert all(p["id"] != policy_id for p in admin_client.get("/api/assistant/policies").json())


def test_ai_012a_employees_see_only_published_policies(admin_client, employee_client):
    """AI-012 — the draft/published split is enforced on the list route."""
    admin_client.post(
        "/api/assistant/policies",
        json={"title": "Draft Only", "category": "General", "content": "Not for employees", "is_published": False},
    )
    admin_client.post(
        "/api/assistant/policies",
        json={"title": "Published One", "category": "General", "content": "For everyone", "is_published": True},
    )

    titles = [p["title"] for p in employee_client.get("/api/assistant/policies").json()]
    assert "Published One" in titles
    assert "Draft Only" not in titles

    admin_titles = [p["title"] for p in admin_client.get("/api/assistant/policies").json()]
    assert "Draft Only" in admin_titles
